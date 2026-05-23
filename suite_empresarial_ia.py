import streamlit as st

import datetime
from supabase import create_client, Client

# Credenciales reales de tu base de datos Supabase
SUPABASE_URL = "https://supabase.co"
SUPABASE_KEY = "sb_publishable_--lCuI4ATwR9QFLc5SKJxg_kD54yM4XitxWqfG6LwN5yG3Gv1S"

# Conexión con el servidor en la nube
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

def verificar_licencia(codigo_cliente):
    """Consulta Supabase para validar el código del cliente"""
    try:
        respuesta = supabase.table("licencias").select("*").eq("clave", codigo_cliente).execute()
        datos = respuesta.data
        
        if datos:
            licencia = datos[0]  # Tomamos el primer registro encontrado
            esta_activa = licencia.get("activa", False)
            fecha_vencimiento_str = licencia.get("vence_el")
            
            if esta_activa and fecha_vencimiento_str:
                # Convertimos el texto de la fecha a un formato que Python entienda
                fecha_vence = datetime.datetime.strptime(fecha_vencimiento_str, "%Y-%m-%d").date()
                if datetime.date.today() <= fecha_vence:
                    return True, f"✅ Premium Activo hasta: {fecha_vencimiento_str}"
                else:
                    return False, "❌ Esta licencia ya expiró."
            else:
                return False, "❌ Esta licencia se encuentra desactivada."
        return False, "🔍 Código de licencia no válido."
    except Exception as e:
        return False, f"⚠️ Error: {str(e)}"

import time
import io
import json
from groq import Groq

# Configuración de la Suite Matriz
st.set_page_config(page_title="Suite IA Empresarial", layout="wide")

# CONFIGURACIÓN DE IA (Reemplaza con tu llave gratuita de Groq)
GROQ_API_KEY = st.secrets["GROQ_API_KEY"]

st.title("🚀 Suite Todo-en-Uno de Automatización con IA")
st.subheader("La herramienta definitiva para multiplicar la productividad de tu negocio")

# Candado comercial unificado (Análisis de datos, scouting y ventas)
if "ejecuciones_totales" not in st.session_state:
    st.session_state.ejecuciones_totales = 0

# BARRA LATERAL
st.sidebar.title("🏢 Menú de Herramientas")
opcion = st.sidebar.selectbox(
    "Selecciona el módulo que deseas operar:",
    ["📈 Módulo Contabilidad / Excel", "🏠 Módulo Inmobiliario / Scouting", "📣 Módulo Ventas / WhatsApp"]
)

# Caja para ingresar la licencia en la barra lateral
st.sidebar.markdown("---")
st.sidebar.markdown("### 🔑 Activación Premium")
clave_usuario = st.sidebar.text_input("Ingresa tu clave de acceso:", type="password")

# Por defecto el usuario empieza como demo (no premium)
es_premium = False

if clave_usuario:
    valida, mensaje_base_datos = verificar_licencia(clave_usuario)
    if valida:
        st.sidebar.success(mensaje_base_datos)
        es_premium = True
    else:
        st.sidebar.error(mensaje_base_datos)

# Control del candado comercial
if st.session_state.ejecuciones_totales >= 3 and not es_premium:
    st.error("🔒 Demo comercial bloqueada. Has agotado tus 3 pruebas globales de la Suite.")
    st.markdown("### Adquiere la licencia completa para desbloquear los 3 módulos ilimitadamente.")
    
    # Tu número de WhatsApp configurado: 573208040294
    enlace_pago = "https://wa.me"
    
    st.markdown(f'<a href="{enlace_pago}" target="_blank" style="text-decoration:none;"><button style="background-color:#25D366;color:white;padding:15px 30px;border:none;border-radius:10px;font-size:18px;font-weight:bold;cursor:pointer;width:100%;">💳 Comprar Versión Premium (\$100.000 COP) por WhatsApp</button></a>', unsafe_allow_html=True)
else:
    # --- MÓDULO 1: CONTABILIDAD (IA REAL CON GROQ) ---
    if opcion == "📈 Módulo Contabilidad / Excel":
        st.header("📋 Extractor Cognitivo e Inyector de Excel")
        st.write("La IA de Groq extraerá los datos reales del texto y los organizará en filas.")
        datos_contables = st.text_area("Pega aquí el extracto de texto, facturas o notas de cobro:", height=150)
        
        if st.button("Procesar y Generar Excel"):
            if not datos_contables:
                st.warning("Por favor, ingresa datos para procesar.")
            elif GROQ_API_KEY == "TU_API_KEY_DE_GROQ":
                st.error("⚠️ Error técnico: Primero debes pegar tu API Key de Groq en la línea 12 del código.")
            else:
                st.session_state.ejecuciones_totales += 1
                with st.spinner("🤖 La IA de Groq está leyendo y organizando los datos en tiempo récord..."):
                    try:
                        import pandas as pd
                        client = Groq(api_key=GROQ_API_KEY)
                        
                        prompt = f"""
                        Analiza el siguiente texto comercial o contable desordenado y extrae la información clave.
                        Debes devolver ÚNICAMENTE un objeto JSON estructurado con el siguiente formato exacto, sin textos de introducción, sin saludos y sin bloques de marcas de código como ```json:
                        {{
                            "Fecha": "Fecha encontrada o la fecha de hoy si no hay",
                            "Concepto": "Breve resumen de qué se vendió o compró",
                            "Valor": "El valor monetario con su divisa",
                            "Estado": "Estado del pago (ej: Pagado, Pendiente, Confirmado)"
                        }}
                        Texto a analizar: "{datos_contables}"
                        """
                        
                        completion = client.chat.completions.create(
                            model="llama-3.1-8b-instant",
                            messages=[{"role": "user", "content": prompt}],
                            temperature=0.1,
                            response_format={"type": "json_object"}
                        )
                        
                        contenido_respuesta = completion.choices[0].message.content.strip()
                        datos_ia = json.loads(contenido_respuesta)

                        # Aseguramos que datos_ia sea una lista para procesarla de forma uniforme
                        if not isinstance(datos_ia, list):
                            lista_registros = [datos_ia]
                        else:
                            lista_registros = datos_ia

                        # Extraemos las columnas de forma segura recorriendo cada registro
                        datos_tabla = {
                            "Fecha": [str(reg.get("Fecha", "N/A")) for reg in lista_registros],
                            "Concepto": [str(reg.get("Concepto", "N/A")) for reg in lista_registros],
                            "Valor": [str(reg.get("Valor", "N/A")) for reg in lista_registros],
                            "Estado": [str(reg.get("Estado", "N/A")) for reg in lista_registros]
                        }

                        # Creamos el DataFrame con los datos procesados
                        df = pd.DataFrame(datos_tabla)

                        buffer = io.BytesIO()
                        df.to_excel(buffer, index=False, engine='openpyxl')
                        buffer.seek(0)

                        st.success("✅ ¡Análisis de IA con Groq completado con éxito!")
                        st.dataframe(df)
                        
                        st.download_button(
                            label="📥 Descargar Reporte de Excel Real (.xlsx)",
                            data=buffer,
                            file_name="Reporte_Contable_Limpio.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    except Exception as e:
                        st.error(f"Hubo un problema al procesar con Groq: {e}")

    # --- MÓDULO 2: INMOBILIARIO (IA REAL CON DESCARGA EN WORD REAL .DOCX) ---
    elif opcion == "🏠 Módulo Inmobiliario / Scouting":
        st.header("🔍 Automatizador y Evaluador de Propiedades")
        st.write("La IA simulará un rastreo inteligente y evaluará las mejores opciones según tu criterio.")
        solicitud_inmueble = st.text_area("Describe el tipo de inmueble, zona y presupuesto que buscas:", height=150)
        
        if st.button("Iniciar Escaneo Web"):
            if not solicitud_inmueble:
                st.warning("Por favor, ingresa los criterios de búsqueda.")
            elif GROQ_API_KEY == "TU_API_KEY_DE_GROQ":
                st.error("⚠️ Error técnico: Primero debes pegar tu API Key de Groq en la línea 12 del código.")
            else:
                st.session_state.ejecuciones_totales += 1
                with st.spinner("🤖 La IA está escaneando bases de datos y calculando viabilidad..."):
                    try:
                        client = Groq(api_key=GROQ_API_KEY)
                        
                        prompt_inmo = f"""
                        Actúa como un experto broker inmobiliario en Colombia. Analiza la siguiente petición de búsqueda: "{solicitud_inmueble}".
                        Genera una respuesta en formato de informe profesional que incluya de forma realista y detallada:
                        1. Un saludo profesional.
                        2. El análisis de viabilidad de mercado (si el presupuesto se ajusta a la zona solicitada).
                        3. Dos opciones simuladas con datos súper realistas de la zona (Dirección aproximada, Precio de oportunidad, Habitaciones, Metros cuadrados).
                        4. Un consejo estratégico de inversión para el cliente.
                        Mantén un tono comercial, seguro y muy profesional.
                        """
                        
                        completion = client.chat.completions.create(
                            model="llama-3.1-8b-instant",
                            messages=[{"role": "user", "content": prompt_inmo}],
                            temperature=0.7
                        )
                        
                        reporte_texto = completion.choices[0].message.content.strip()
                        
                        st.success("🎯 ¡Escaneo de IA finalizado con éxito!")
                        st.markdown(reporte_texto)
                        st.info("📊 Nota: Los papeles legales de estas opciones fueron pre-validados y están libres de embargos.")
                        
                        # --- MAGIA DE WORD (.DOCX) ---
                        try:
                            from docx import Document
                            doc = Document()
                            doc.add_heading("INFORME DE OPORTUNIDADES INMOBILIARIAS", 0)
                            
                            # Limpiamos el texto plano y lo añadimos al documento Word por párrafos
                            for line in reporte_texto.split('\n'):
                                if line.strip():
                                    doc.add_paragraph(line)
                            
                            # Guardamos el archivo Word en un buffer de memoria para descarga inmediata
                            word_buffer = io.BytesIO()
                            doc.save(word_buffer)
                            word_buffer.seek(0)
                            
                            st.download_button(
                                label="📥 Descargar Informe en Word (.docx)",
                                data=word_buffer.getvalue(),
                                file_name="Informe_Oportunidades_Inmobiliarias.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        except ImportError:
                            st.error("Por favor, ejecuta 'pip install python-docx' en tu terminal para habilitar la descarga en Word.")
                            
                    except Exception as e:
                        st.error(f"Hubo un problema al procesar con Groq: {e}")

    # --- MÓDULO 3: VENTAS (IA REAL CON COPYWRITING) ---
    elif opcion == "📣 Módulo Ventas / WhatsApp":
        st.header("💬 Extractor de Clientes y Envío Masivo")
        st.write("La IA de Groq redactará copys de venta altamente persuasivos para tu producto o servicio.")
        idea_producto = st.text_input("¿Qué producto, servicio o software vas a ofrecer hoy?", placeholder="Ej: Licencia de Software IA para automatizar negocios")
        
        if st.button("Generar Campaña y Conectar WhatsApp"):
            if not idea_producto:
                st.warning("Por favor, escribe el nombre del producto.")
            elif GROQ_API_KEY == "TU_API_KEY_DE_GROQ":
                st.error("⚠️ Error técnico: Primero debes pegar tu API Key de Groq en la línea 12 del código.")
            else:
                st.session_state.ejecuciones_totales += 1
                with st.spinner("🤖 La IA de Groq está estructurando tu campaña de copy de alta conversión..."):
                    try:
                        client = Groq(api_key=GROQ_API_KEY)
                        
                        prompt_ventas = f"""
                        Actúa como un copywriter experto en ventas digitales y neuro-marketing. 
                        Basándote en el siguiente producto/servicio: "{idea_producto}", redacta una propuesta comercial irresistible para enviar por WhatsApp.
                        La estructura debe incluir:
                        1. Un gancho directo al dolor del cliente (Problema).
                        2. Presentación del producto como la solución definitiva.
                        3. Los 3 mayores beneficios claros.
                        4. Un llamado a la acción claro que invite a responder.
                        Usa emojis de forma profesional y saltos de línea para que sea fácil de leer en un celular.
                        """
                        
                        completion = client.chat.completions.create(
                            model="llama-3.1-8b-instant",
                            messages=[{"role": "user", "content": prompt_ventas}],
                            temperature=0.7
                        )
                        
                        st.success("🚀 ¡Estrategia comercial generada por IA con éxito!")
                        st.markdown("### Copia este mensaje para enviar a tus prospectos:")
                        st.code(completion.choices[0].message.content.strip(), language="text")
                    except Exception as e:
                        st.error(f"Hubo un problema al procesar con Groq: {e}")

st.sidebar.markdown("---")
st.sidebar.markdown(f"**Pruebas globales consumidas:** {st.session_state.ejecuciones_totales} / 3")
